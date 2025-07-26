#!/usr/bin/env python3
"""
报告导出服务
提供医疗AI预测结果的报告导出功能
支持PDF、Excel、Word等多种格式
"""

import os
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Union

# 导出格式处理
import pandas as pd
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class ReportExportService:
    """报告导出服务"""
    
    def __init__(self, base_dir: Path):
        """
        初始化报告导出服务
        
        Args:
            base_dir: 项目根目录
        """
        self.base_dir = base_dir
        self.export_dir = base_dir / "exports"
        self.template_dir = base_dir / "backend" / "templates" / "reports"
        
        # 创建导出目录
        self.export_dir.mkdir(exist_ok=True, parents=True)
        logger.info(f"导出目录: {self.export_dir}")
        
    def export_report(self, 
                     record_data: Dict[str, Any],
                     export_format: str = 'pdf',
                     template_name: Optional[str] = None) -> Dict[str, Any]:
        """
        导出报告
        
        Args:
            record_data: 记录数据
            export_format: 导出格式 (pdf, excel, word)
            template_name: 模板名称
            
        Returns:
            包含导出结果的字典
        """
        try:
            # 根据格式选择导出方法
            if export_format.lower() == 'pdf':
                return self._export_pdf(record_data, template_name)
            elif export_format.lower() == 'excel':
                return self._export_excel(record_data)
            elif export_format.lower() == 'word':
                return self._export_word(record_data, template_name)
            else:
                raise ValueError(f"不支持的导出格式: {export_format}")
        except Exception as e:
            logger.error(f"导出报告失败: {e}")
            raise
            
    def _translate_recommendation(self, text):
        """将中文健康建议转换为英文"""
        translations = {
            "定期进行体检，监测身体状况": "Regular health checkups to monitor physical condition",
            "保持健康的生活方式，适量运动": "Maintain a healthy lifestyle with moderate exercise",
            "均衡饮食，避免高糖高脂食物": "Balanced diet, avoid high-sugar and high-fat foods",
            "如有不适症状，请及时就医": "Seek medical attention promptly if you experience discomfort",
            "建议咨询专业医生获取更详细的医疗建议": "Consult a professional doctor for more detailed medical advice",
            "建议立即就医进行详细检查": "Immediate medical examination is recommended",
            "定期监测血压、血糖和胆固醇": "Regularly monitor blood pressure, blood sugar and cholesterol",
            "避免剧烈运动，保持情绪稳定": "Avoid strenuous exercise, maintain emotional stability",
            "戒烟限酒，控制饮食": "Quit smoking, limit alcohol, control diet",
            "建议立即就医进行进一步检查": "Immediate further medical examination is recommended",
            "定期进行相关筛查": "Regular relevant screenings",
            "避免吸烟和有害物质暴露": "Avoid smoking and exposure to harmful substances",
            "保持积极心态，配合治疗": "Maintain a positive attitude and cooperate with treatment",
            "定期监测血糖和并发症相关指标": "Regularly monitor blood sugar and complication-related indicators",
            "严格控制饮食，减少糖分和碳水化合物摄入": "Strictly control diet, reduce sugar and carbohydrate intake",
            "适量运动，保持健康体重": "Exercise moderately, maintain healthy weight",
            "定期进行并发症筛查": "Regular screening for complications",
            "建议就医进行糖尿病筛查": "Medical diabetes screening is recommended",
            "定期监测血糖水平": "Regularly monitor blood sugar levels",
            "控制饮食，减少糖分摄入": "Control diet, reduce sugar intake",
            "注意并发症早期症状": "Pay attention to early symptoms of complications",
            "继续保持当前健康管理方案": "Continue with current health management plan",
            "建议咨询专业医生": "Consultation with a professional doctor is recommended",
            "定期监测相关指标": "Regularly monitor relevant indicators",
            "定期进行胸部X光检查": "Regular chest X-ray examinations",
            "注意呼吸系统健康": "Pay attention to respiratory system health"
        }
        
        # 如果找到完全匹配的翻译，则返回
        if text in translations:
            return translations[text]
            
        # 如果没有完全匹配，尝试部分匹配
        for zh, en in translations.items():
            if zh in text:
                return en
                
        # 如果没有匹配，返回原文
        return text
        
    def _safe_text(self, text):
        """
        安全处理文本，将非ASCII字符替换为空格
        这样可以保留大部分可读字符，同时避免编码问题
        """
        if not isinstance(text, str):
            text = str(text)
        
        # 将非ASCII字符替换为空格，而不是完全移除
        return ''.join([c if ord(c) < 128 else ' ' for c in text])

    def _export_pdf(self, record_data: Dict[str, Any], template_name: Optional[str] = None) -> Dict[str, Any]:
        """
        导出PDF报告
        
        Args:
            record_data: 记录数据
            template_name: 模板名称
            
        Returns:
            包含导出结果的字典
        """
        try:
            # 创建PDF对象
            pdf = FPDF()
            pdf.add_page()
            
            # 设置中文字体（使用微软雅黑字体）
            msyh_font_path = str(self.base_dir / "backend" / "fonts" / "msyh.ttc")
            if not os.path.exists(msyh_font_path):
                # 如果找不到微软雅黑字体，尝试使用其他可能的路径
                msyh_font_path = str(Path(__file__).parent.parent / "fonts" / "msyh.ttc")
                if not os.path.exists(msyh_font_path):
                    # 如果还是找不到，使用系统默认字体
                    msyh_font_path = None
            
            # 添加字体并设置默认字体
            try:
                pdf.set_font('Arial', '', 12)
                logger.info("使用默认字体Arial")
            except Exception as font_error:
                # 如果加载字体失败，使用默认字体并记录错误
                logger.error(f"加载字体失败: {font_error}，使用默认字体")
                pdf.set_font('Arial', '', 12)
            
            # 添加标题
            pdf.set_font_size(18)
            pdf.cell(200, 10, self._safe_text('Medical AI Diagnosis Report'), 0, 1, 'C')
            pdf.ln(10)
            
            # 添加基本信息
            pdf.set_font_size(12)
            pdf.cell(200, 10, self._safe_text(f"Report Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"), 0, 1)
            pdf.cell(200, 10, self._safe_text(f"Model Type: {self._get_model_display_name(record_data.get('model_type', ''))}"), 0, 1)
            pdf.cell(200, 10, self._safe_text(f"Risk Level: {self._get_risk_text(record_data.get('risk_level', ''))}"), 0, 1)
            pdf.cell(200, 10, self._safe_text(f"Prediction Time: {record_data.get('created_at_local', '')}"), 0, 1)
            pdf.ln(10)
            
            # 添加预测摘要
            pdf.set_font_size(14)
            pdf.cell(200, 10, self._safe_text('Prediction Summary'), 0, 1)
            pdf.set_font_size(12)
            summary = record_data.get('summary', 'No summary available')
            # 处理中文字符，转换为英文或移除
            if isinstance(summary, str):
                summary = summary.replace('预测完成', 'Prediction completed')
                summary = summary.replace('未检测到明显疾病', 'No obvious diseases detected')
                summary = summary.replace('检测到疾病', 'Diseases detected')
            pdf.multi_cell(0, 10, self._safe_text(summary))
            pdf.ln(10)
            
            # 添加详细结果
            pdf.set_font_size(14)
            pdf.cell(200, 10, self._safe_text('Detailed Results'), 0, 1)
            pdf.set_font_size(12)
            
            # 根据不同模型类型显示不同内容
            model_type = record_data.get('model_type', '')
            prediction_result = record_data.get('prediction_result', {})
            
            if model_type == 'medical_qa':
                pdf.multi_cell(0, 10, self._safe_text(f"Question: {prediction_result.get('question', '')}"))
                pdf.multi_cell(0, 10, self._safe_text(f"Answer: {prediction_result.get('answer', '')}"))
                
            elif model_type == 'heart_disease':
                pdf.multi_cell(0, 10, self._safe_text(f"Prediction: {prediction_result.get('prediction', '')}"))
                pdf.multi_cell(0, 10, self._safe_text(f"Confidence: {prediction_result.get('confidence', 0) * 100:.1f}%"))
                
            elif model_type == 'tumor':
                pdf.multi_cell(0, 10, self._safe_text(f"Tumor Type: {prediction_result.get('prediction', '')}"))
                pdf.multi_cell(0, 10, self._safe_text(f"Confidence: {prediction_result.get('confidence', 0) * 100:.1f}%"))
                
            elif model_type == 'diabetes':
                pdf.multi_cell(0, 10, self._safe_text(f"Diabetes Risk: {prediction_result.get('prediction', '')}"))
                pdf.multi_cell(0, 10, self._safe_text(f"Complication Type: {prediction_result.get('并发症类型', 'None')}"))
                pdf.multi_cell(0, 10, self._safe_text(f"Risk Probability: {prediction_result.get('发病概率', '0%')}"))
                
            elif model_type == 'chest_xray':
                # 胸部X光特殊处理
                predictions = prediction_result.get('predictions', {})
                pdf.multi_cell(0, 10, self._safe_text("Detected Diseases:"))
                
                for disease, data in predictions.items():
                    if data.get('positive', False):
                        pdf.multi_cell(0, 10, self._safe_text(f"- {disease}: {data.get('probability', 0) * 100:.1f}%"))
                
                # 如果有热力图，添加图片路径信息
                if 'heatmap_path' in prediction_result:
                    pdf.multi_cell(0, 10, self._safe_text("Heatmap generated, view in system"))
            
            pdf.ln(10)
            
            # 添加健康建议
            pdf.set_font_size(14)
            pdf.cell(200, 10, self._safe_text('Health Recommendations'), 0, 1)
            pdf.set_font_size(12)
            
            recommendations = record_data.get('recommendations', [])
            for i, recommendation in enumerate(recommendations):
                # 先翻译中文建议为英文，再过滤非法字符
                translated_recommendation = self._translate_recommendation(recommendation)
                pdf.multi_cell(0, 10, self._safe_text(f"{i+1}. {translated_recommendation}"))
            
            # 添加页脚
            pdf.set_y(-30)
            pdf.set_font_size(10)
            pdf.cell(0, 10, self._safe_text('This report is automatically generated by Medical AI Diagnosis Platform, for reference only'), 0, 1, 'C')
            pdf.cell(0, 10, self._safe_text('Please consult a professional doctor if you have any questions'), 0, 1, 'C')
            
            # 生成文件名
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            filename = f"{model_type}_{timestamp}.pdf"
            filepath = self.export_dir / filename
            
            # 保存PDF
            pdf.output(str(filepath))
            
            return {
                'success': True,
                'message': '报告导出成功',
                'filename': filename,
                'filepath': str(filepath),
                'format': 'pdf'
            }
            
        except Exception as e:
            logger.error(f"导出PDF报告失败: {e}")
            return {
                'success': False,
                'message': f'报告导出失败: {str(e)}',
                'format': 'pdf'
            }
            
    def _export_excel(self, record_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        导出Excel报告
        
        Args:
            record_data: 记录数据
            
        Returns:
            包含导出结果的字典
        """
        try:
            # 创建数据字典
            report_data = {
                'Report Information': [],
                'Prediction Results': [],
                'Health Recommendations': []
            }
            
            # 添加基本信息
            report_data['Report Information'] = [
                ['Report Generation Time', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
                ['Model Type', self._get_model_display_name(record_data.get('model_type', ''))],
                ['Risk Level', self._get_risk_text(record_data.get('risk_level', ''))],
                ['Prediction Time', record_data.get('created_at_local', '')],
                ['Prediction Summary', record_data.get('summary', 'No summary available')]
            ]
            
            # 根据不同模型类型添加不同的预测结果
            model_type = record_data.get('model_type', '')
            prediction_result = record_data.get('prediction_result', {})
            
            if model_type == 'medical_qa':
                report_data['Prediction Results'] = [
                    ['Question', prediction_result.get('question', '')],
                    ['Answer', prediction_result.get('answer', '')]
                ]
                
            elif model_type == 'heart_disease':
                report_data['Prediction Results'] = [
                    ['Prediction', prediction_result.get('prediction', '')],
                    ['Confidence', f"{prediction_result.get('confidence', 0) * 100:.1f}%"]
                ]
                
                # 添加输入特征
                input_data = record_data.get('input_data', {})
                for key, value in input_data.items():
                    report_data['Prediction Results'].append([self._get_data_label(key), value])
                
            elif model_type == 'tumor':
                report_data['Prediction Results'] = [
                    ['Tumor Type', prediction_result.get('prediction', '')],
                    ['Confidence', f"{prediction_result.get('confidence', 0) * 100:.1f}%"]
                ]
                
            elif model_type == 'diabetes':
                report_data['Prediction Results'] = [
                    ['Diabetes Risk', prediction_result.get('prediction', '')],
                    ['Complication Type', prediction_result.get('并发症类型', 'None')],
                    ['Risk Probability', prediction_result.get('发病概率', '0%')]
                ]
                
                # 添加输入特征
                input_data = record_data.get('input_data', {})
                for key, value in input_data.items():
                    report_data['Prediction Results'].append([self._get_data_label(key), value])
                
            elif model_type == 'chest_xray':
                # 胸部X光特殊处理
                predictions = prediction_result.get('predictions', {})
                
                # 添加所有预测结果
                result_rows = []
                for disease, data in predictions.items():
                    result_rows.append([
                        disease, 
                        'Positive' if data.get('positive', False) else 'Negative',
                        f"{data.get('probability', 0) * 100:.1f}%"
                    ])
                
                # 如果有结果，添加表头
                if result_rows:
                    report_data['Prediction Results'].append(['Disease', 'Result', 'Probability'])
                    report_data['Prediction Results'].extend(result_rows)
                
                # 如果有热力图，添加图片路径信息
                if 'heatmap_path' in prediction_result:
                    report_data['Prediction Results'].append(['Heatmap', 'Generated, view in system'])
            
            # 添加健康建议
            recommendations = record_data.get('recommendations', [])
            for i, recommendation in enumerate(recommendations):
                # 先翻译中文建议为英文
                translated_recommendation = self._translate_recommendation(recommendation)
                report_data['Health Recommendations'].append([f'Recommendation {i+1}', translated_recommendation])
            
            # 创建Excel文件
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            filename = f"{model_type}_{timestamp}.xlsx"
            filepath = self.export_dir / filename
            
            # 创建Excel写入器
            with pd.ExcelWriter(filepath) as writer:
                # 写入报告信息
                pd.DataFrame(report_data['Report Information'], columns=['Item', 'Content']).to_excel(
                    writer, sheet_name='Report Information', index=False
                )
                
                # 写入预测结果
                if model_type == 'chest_xray' and report_data['Prediction Results'] and len(report_data['Prediction Results'][0]) == 3:
                    # 胸部X光特殊处理
                    pd.DataFrame(report_data['Prediction Results'][1:], columns=report_data['Prediction Results'][0]).to_excel(
                        writer, sheet_name='Prediction Results', index=False
                    )
                else:
                    pd.DataFrame(report_data['Prediction Results'], columns=['Item', 'Content']).to_excel(
                        writer, sheet_name='Prediction Results', index=False
                    )
                
                # 写入健康建议
                pd.DataFrame(report_data['Health Recommendations'], columns=['No.', 'Recommendation']).to_excel(
                    writer, sheet_name='Health Recommendations', index=False
                )
            
            return {
                'success': True,
                'message': '报告导出成功',
                'filename': filename,
                'filepath': str(filepath),
                'format': 'excel'
            }
            
        except Exception as e:
            logger.error(f"导出Excel报告失败: {e}")
            return {
                'success': False,
                'message': f'报告导出失败: {str(e)}',
                'format': 'excel'
            }
            
    def _export_word(self, record_data: Dict[str, Any], template_name: Optional[str] = None) -> Dict[str, Any]:
        """
        导出Word报告
        
        Args:
            record_data: 记录数据
            template_name: 模板名称
            
        Returns:
            包含导出结果的字典
        """
        try:
            # 创建Word文档
            doc = Document()
            
            # 添加标题
            title = doc.add_heading('Medical AI Diagnosis Report', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加基本信息
            doc.add_heading('Basic Information', level=1)
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            
            # 填充基本信息表格
            cells = table.rows[0].cells
            cells[0].text = 'Report Generation Time'
            cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            cells = table.rows[1].cells
            cells[0].text = 'Model Type'
            cells[1].text = self._get_model_display_name(record_data.get('model_type', ''))
            
            cells = table.rows[2].cells
            cells[0].text = 'Risk Level'
            cells[1].text = self._get_risk_text(record_data.get('risk_level', ''))
            
            cells = table.rows[3].cells
            cells[0].text = 'Prediction Time'
            cells[1].text = record_data.get('created_at_local', '')
            
            doc.add_paragraph()
            
            # 添加预测摘要
            doc.add_heading('Prediction Summary', level=1)
            summary = record_data.get('summary', 'No summary available')
            # 处理中文字符
            if isinstance(summary, str):
                summary = summary.replace('预测完成', 'Prediction completed')
                summary = summary.replace('未检测到明显疾病', 'No obvious diseases detected')
                summary = summary.replace('检测到疾病', 'Diseases detected')
            doc.add_paragraph(summary)
            
            # 添加详细结果
            doc.add_heading('Detailed Results', level=1)
            
            # 根据不同模型类型显示不同内容
            model_type = record_data.get('model_type', '')
            prediction_result = record_data.get('prediction_result', {})
            
            if model_type == 'medical_qa':
                p = doc.add_paragraph()
                p.add_run('Question: ').bold = True
                p.add_run(prediction_result.get('question', ''))
                
                p = doc.add_paragraph()
                p.add_run('Answer: ').bold = True
                p.add_run(prediction_result.get('answer', ''))
                
            elif model_type == 'heart_disease':
                p = doc.add_paragraph()
                p.add_run('Prediction: ').bold = True
                p.add_run(prediction_result.get('prediction', ''))
                
                p = doc.add_paragraph()
                p.add_run('Confidence: ').bold = True
                p.add_run(f"{prediction_result.get('confidence', 0) * 100:.1f}%")
                
                # 添加输入特征表格
                doc.add_heading('Input Features', level=2)
                input_data = record_data.get('input_data', {})
                if input_data:
                    table = doc.add_table(rows=len(input_data), cols=2)
                    table.style = 'Table Grid'
                    
                    for i, (key, value) in enumerate(input_data.items()):
                        cells = table.rows[i].cells
                        cells[0].text = self._get_data_label(key)
                        cells[1].text = str(value)
                
            elif model_type == 'tumor':
                p = doc.add_paragraph()
                p.add_run('Tumor Type: ').bold = True
                p.add_run(prediction_result.get('prediction', ''))
                
                p = doc.add_paragraph()
                p.add_run('Confidence: ').bold = True
                p.add_run(f"{prediction_result.get('confidence', 0) * 100:.1f}%")
                
            elif model_type == 'diabetes':
                p = doc.add_paragraph()
                p.add_run('Diabetes Risk: ').bold = True
                p.add_run(prediction_result.get('prediction', ''))
                
                p = doc.add_paragraph()
                p.add_run('Complication Type: ').bold = True
                p.add_run(prediction_result.get('并发症类型', 'None'))
                
                p = doc.add_paragraph()
                p.add_run('Risk Probability: ').bold = True
                p.add_run(prediction_result.get('发病概率', '0%'))
                
                # 添加输入特征表格
                doc.add_heading('Input Features', level=2)
                input_data = record_data.get('input_data', {})
                if input_data:
                    table = doc.add_table(rows=len(input_data), cols=2)
                    table.style = 'Table Grid'
                    
                    for i, (key, value) in enumerate(input_data.items()):
                        cells = table.rows[i].cells
                        cells[0].text = self._get_data_label(key)
                        cells[1].text = str(value)
                
            elif model_type == 'chest_xray':
                # 胸部X光特殊处理
                predictions = prediction_result.get('predictions', {})
                
                doc.add_paragraph('Detected Diseases:')
                
                # 创建疾病表格
                positive_diseases = [
                    (disease, data) 
                    for disease, data in predictions.items() 
                    if data.get('positive', False)
                ]
                
                if positive_diseases:
                    table = doc.add_table(rows=len(positive_diseases) + 1, cols=3)
                    table.style = 'Table Grid'
                    
                    # 添加表头
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'Disease'
                    header_cells[1].text = 'Result'
                    header_cells[2].text = 'Probability'
                    
                    # 添加疾病数据
                    for i, (disease, data) in enumerate(positive_diseases):
                        cells = table.rows[i + 1].cells
                        cells[0].text = disease
                        cells[1].text = 'Positive'
                        cells[2].text = f"{data.get('probability', 0) * 100:.1f}%"
                else:
                    doc.add_paragraph('No obvious diseases detected')
                
                # 如果有热力图，添加图片路径信息
                if 'heatmap_path' in prediction_result:
                    doc.add_paragraph('Heatmap generated, view in system')
            
            # 添加健康建议
            doc.add_heading('Health Recommendations', level=1)
            recommendations = record_data.get('recommendations', [])
            
            if recommendations:
                for i, recommendation in enumerate(recommendations):
                    # 先翻译中文建议为英文
                    translated_recommendation = self._translate_recommendation(recommendation)
                    doc.add_paragraph(f"{i+1}. {translated_recommendation}", style='List Number')
            else:
                doc.add_paragraph('No health recommendations available')
            
            # 添加页脚
            footer_section = doc.sections[0]
            footer = footer_section.footer
            footer_text = footer.paragraphs[0]
            footer_text.text = 'This report is automatically generated by Medical AI Diagnosis Platform, for reference only. Please consult a professional doctor if you have any questions.'
            footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 生成文件名
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            filename = f"{model_type}_{timestamp}.docx"
            filepath = self.export_dir / filename
            
            # 保存Word文档
            doc.save(filepath)
            
            return {
                'success': True,
                'message': '报告导出成功',
                'filename': filename,
                'filepath': str(filepath),
                'format': 'word'
            }
            
        except Exception as e:
            logger.error(f"导出Word报告失败: {e}")
            return {
                'success': False,
                'message': f'报告导出失败: {str(e)}',
                'format': 'word'
            }
            
    def _get_model_display_name(self, model_type: str) -> str:
        """获取模型显示名称"""
        model_names = {
            'medical_qa': 'Medical Q&A',
            'heart_disease': 'Heart Disease Prediction',
            'tumor': 'Tumor Classification',
            'diabetes': 'Diabetes Assessment',
            'chest_xray': 'Chest X-ray Detection'
        }
        return model_names.get(model_type, model_type)
        
    def _get_risk_text(self, risk_level: str) -> str:
        """获取风险等级文本"""
        risk_texts = {
            'high': 'High Risk',
            'medium': 'Medium Risk',
            'low': 'Low Risk',
            'info': 'Information'
        }
        return risk_texts.get(risk_level, 'Unknown')
        
    def _get_data_label(self, key: str) -> str:
        """获取数据标签"""
        labels = {
            'age': 'Age',
            'sex': 'Sex',
            'gender': 'Gender',
            'cp': 'Chest Pain Type',
            'trestbps': 'Resting Blood Pressure',
            'chol': 'Cholesterol',
            'fbs': 'Fasting Blood Sugar',
            'restecg': 'ECG Results',
            'thalach': 'Max Heart Rate',
            'exang': 'Exercise Induced Angina',
            'oldpeak': 'ST Depression',
            'slope': 'ST Slope',
            'ca': 'Number of Major Vessels',
            'thal': 'Thalassemia',
            'pregnancies': 'Pregnancies',
            'glucose': 'Glucose Level',
            'blood_pressure': 'Blood Pressure',
            'skin_thickness': 'Skin Thickness',
            'insulin': 'Insulin',
            'bmi': 'BMI',
            'diabetes_pedigree': 'Diabetes Family History',
            'smoking': 'Smoking History',
            'yellow_fingers': 'Yellow Fingers',
            'anxiety': 'Anxiety',
            'peer_pressure': 'Peer Pressure',
            'chronic_disease': 'Chronic Disease',
            'fatigue': 'Fatigue',
            'allergy': 'Allergy',
            'wheezing': 'Wheezing',
            'alcohol_consuming': 'Alcohol Consumption',
            'coughing': 'Coughing',
            'shortness_of_breath': 'Shortness of Breath',
            'swallowing_difficulty': 'Swallowing Difficulty',
            'chest_pain': 'Chest Pain',
            'lung_cancer': 'Lung Cancer Family History'
        }
        return labels.get(key, key) 